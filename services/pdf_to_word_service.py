"""
PDF to Word Service for converting PDF documents to Word (.docx) format.
Integrates OCR capabilities for scanned/image-based PDFs.
"""
import os
import tempfile
import shutil
from pathlib import Path
from typing import Optional, Callable, List, Tuple
from dataclasses import dataclass
from enum import Enum

import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from services.ocr_service import OCRService, AccuracyMode


class ConversionMode(Enum):
    """Conversion mode options."""
    AUTO = "auto"           # Auto-detect: use OCR if needed
    TEXT_ONLY = "text"      # Extract text only (faster)
    OCR_ALWAYS = "ocr"      # Always use OCR (for scanned PDFs)
    PRESERVE_LAYOUT = "layout"  # Try to preserve layout with images


@dataclass
class PDFToWordSettings:
    """Settings for PDF to Word conversion."""
    conversion_mode: ConversionMode = ConversionMode.AUTO
    language: str = "eng"
    dpi: int = 300
    accuracy_mode: AccuracyMode = AccuracyMode.BALANCED
    include_images: bool = True  # Include images from PDF
    preserve_formatting: bool = True  # Try to preserve bold, italic, etc.


@dataclass
class ConversionResult:
    """Result of PDF to Word conversion."""
    success: bool
    output_path: Optional[str] = None
    total_pages: int = 0
    pages_converted: int = 0
    used_ocr: bool = False
    error_message: Optional[str] = None


class PDFToWordService:
    """
    Service for converting PDF documents to Word format.
    
    Features:
    - Text extraction from digital PDFs
    - OCR integration for scanned PDFs
    - Image extraction and embedding
    - Basic formatting preservation
    """
    
    def __init__(self):
        """Initialize the PDF to Word service."""
        self.ocr_service = OCRService()
    
    def is_tesseract_available(self) -> bool:
        """Check if Tesseract OCR is available."""
        return self.ocr_service.is_tesseract_available()
    
    def _pdf_page_to_image(self, page: fitz.Page, dpi: int = 300) -> Image.Image:
        """
        Convert a PDF page to a PIL Image using PyMuPDF.
        This avoids the need for external Poppler installation.
        
        Args:
            page: PyMuPDF page object.
            dpi: Resolution for rendering.
            
        Returns:
            PIL Image of the page.
        """
        # Calculate zoom factor based on DPI (72 is the base DPI for PDF)
        zoom = dpi / 72.0
        matrix = fitz.Matrix(zoom, zoom)
        
        # Render page to pixmap
        pixmap = page.get_pixmap(matrix=matrix, alpha=False)
        
        # Convert to PIL Image
        img = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
        
        return img
    
    def pdf_has_text(self, pdf_path: str) -> Tuple[bool, int]:
        """
        Check if a PDF contains extractable text.
        
        Args:
            pdf_path: Path to the PDF file.
            
        Returns:
            Tuple of (has_text, page_count)
        """
        try:
            doc = fitz.open(pdf_path)
            page_count = len(doc)
            
            # Check first few pages for text
            pages_to_check = min(3, page_count)
            for i in range(pages_to_check):
                page = doc[i]
                text = page.get_text("text")
                if text and text.strip():
                    doc.close()
                    return True, page_count
            
            doc.close()
            return False, page_count
        except Exception:
            return False, 0
    
    def convert(
        self,
        pdf_path: str,
        output_path: str,
        settings: PDFToWordSettings,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> ConversionResult:
        """
        Convert a PDF file to Word document.
        
        Args:
            pdf_path: Path to the input PDF.
            output_path: Path for the output Word document.
            settings: Conversion settings.
            progress_callback: Optional callback(current_page, total_pages, status_message)
            
        Returns:
            ConversionResult with conversion outcome.
        """
        try:
            has_text, page_count = self.pdf_has_text(pdf_path)
            
            # Determine if we need OCR
            use_ocr = False
            if settings.conversion_mode == ConversionMode.OCR_ALWAYS:
                use_ocr = True
            elif settings.conversion_mode == ConversionMode.AUTO:
                use_ocr = not has_text
            elif settings.conversion_mode == ConversionMode.TEXT_ONLY:
                use_ocr = False
            elif settings.conversion_mode == ConversionMode.PRESERVE_LAYOUT:
                use_ocr = not has_text
            
            if use_ocr and not self.is_tesseract_available():
                return ConversionResult(
                    success=False,
                    error_message="OCR required but Tesseract is not installed."
                )
            
            if progress_callback:
                mode_str = "with OCR" if use_ocr else "extracting text"
                progress_callback(0, page_count, f"Converting PDF {mode_str}...")
            
            if settings.conversion_mode == ConversionMode.PRESERVE_LAYOUT:
                return self._convert_preserve_layout(
                    pdf_path, output_path, settings, use_ocr, progress_callback
                )
            else:
                return self._convert_text_based(
                    pdf_path, output_path, settings, use_ocr, progress_callback
                )
                
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"Conversion failed: {str(e)}"
            )
    
    def _convert_text_based(
        self,
        pdf_path: str,
        output_path: str,
        settings: PDFToWordSettings,
        use_ocr: bool,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> ConversionResult:
        """Convert PDF to Word with text extraction (with or without OCR)."""
        temp_dir = None
        try:
            doc = fitz.open(pdf_path)
            page_count = len(doc)
            
            # Create Word document
            word_doc = Document()
            
            # Set up styles
            self._setup_styles(word_doc)
            
            pages_converted = 0
            
            if use_ocr:
                # Use OCR for text extraction (using PyMuPDF for page rendering)
                import pytesseract
                
                temp_dir = tempfile.mkdtemp()
                
                if progress_callback:
                    progress_callback(0, page_count, "Converting PDF pages to images...")
                
                for i in range(page_count):
                    if progress_callback:
                        progress_callback(i + 1, page_count, f"OCR processing page {i + 1}...")
                    
                    # Convert page to image using PyMuPDF
                    image = self._pdf_page_to_image(doc[i], settings.dpi)
                    
                    # Preprocess image for better OCR
                    processed_image = self.ocr_service._preprocess_image(
                        image, settings.accuracy_mode
                    )
                    
                    # Get OCR text with formatting hints
                    ocr_data = pytesseract.image_to_data(
                        processed_image,
                        lang=settings.language,
                        config=self.ocr_service._get_tesseract_config(settings.accuracy_mode),
                        output_type=pytesseract.Output.DICT
                    )
                    
                    # Process OCR data into paragraphs
                    self._add_ocr_text_to_doc(word_doc, ocr_data, settings)
                    
                    # Add page break except for last page
                    if i < page_count - 1:
                        word_doc.add_page_break()
                    
                    pages_converted += 1
                    
                    # Also extract images if requested
                    if settings.include_images:
                        self._extract_images_from_page(doc[i], word_doc, temp_dir)
            else:
                # Extract text directly from PDF
                for i in range(page_count):
                    if progress_callback:
                        progress_callback(i + 1, page_count, f"Extracting text from page {i + 1}...")
                    
                    page = doc[i]
                    
                    # Extract text blocks with formatting
                    blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                    
                    for block in blocks:
                        if block["type"] == 0:  # Text block
                            self._add_text_block_to_doc(word_doc, block, settings)
                        elif block["type"] == 1 and settings.include_images:  # Image block
                            if temp_dir is None:
                                temp_dir = tempfile.mkdtemp()
                            self._add_image_block_to_doc(word_doc, block, temp_dir)
                    
                    # Add page break except for last page
                    if i < page_count - 1:
                        word_doc.add_page_break()
                    
                    pages_converted += 1
            
            doc.close()
            
            # Save Word document
            word_doc.save(output_path)
            
            return ConversionResult(
                success=True,
                output_path=output_path,
                total_pages=page_count,
                pages_converted=pages_converted,
                used_ocr=use_ocr
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"Conversion failed: {str(e)}"
            )
        finally:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
    
    def _convert_preserve_layout(
        self,
        pdf_path: str,
        output_path: str,
        settings: PDFToWordSettings,
        use_ocr: bool,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> ConversionResult:
        """
        Convert PDF to Word preserving layout by rendering pages as images
        with text overlay.
        """
        temp_dir = None
        try:
            doc = fitz.open(pdf_path)
            page_count = len(doc)
            
            # Create Word document
            word_doc = Document()
            
            # Set margins to minimal for better layout
            for section in word_doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            temp_dir = tempfile.mkdtemp()
            
            if progress_callback:
                progress_callback(0, page_count, "Rendering PDF pages...")
            
            pages_converted = 0
            
            for i in range(page_count):
                if progress_callback:
                    status = f"Processing page {i + 1}..."
                    if use_ocr:
                        status = f"OCR processing page {i + 1}..."
                    progress_callback(i + 1, page_count, status)
                
                # Convert page to image using PyMuPDF
                image = self._pdf_page_to_image(doc[i], settings.dpi)
                
                # Save page image
                img_path = os.path.join(temp_dir, f"page_{i + 1}.png")
                image.save(img_path, "PNG")
                
                # Add image to Word document
                # Calculate size to fit page
                page_width = word_doc.sections[0].page_width - \
                             word_doc.sections[0].left_margin - \
                             word_doc.sections[0].right_margin
                
                word_doc.add_picture(img_path, width=page_width)
                
                # If OCR is needed, add text as invisible/small overlay
                # (for searchability)
                if use_ocr:
                    import pytesseract
                    
                    processed_image = self.ocr_service._preprocess_image(
                        image, settings.accuracy_mode
                    )
                    
                    text = pytesseract.image_to_string(
                        processed_image,
                        lang=settings.language,
                        config=self.ocr_service._get_tesseract_config(settings.accuracy_mode)
                    )
                    
                    if text and text.strip():
                        # Add hidden text paragraph for searchability
                        para = word_doc.add_paragraph()
                        run = para.add_run(text.strip())
                        run.font.size = Pt(1)
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White (invisible)
                
                # Add page break except for last page
                if i < page_count - 1:
                    word_doc.add_page_break()
                
                pages_converted += 1
            
            doc.close()
            
            # Save Word document
            word_doc.save(output_path)
            
            return ConversionResult(
                success=True,
                output_path=output_path,
                total_pages=page_count,
                pages_converted=pages_converted,
                used_ocr=use_ocr
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"Layout conversion failed: {str(e)}"
            )
        finally:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
    
    def _setup_styles(self, word_doc: Document):
        """Set up document styles."""
        # Set default font
        style = word_doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
    def _add_text_block_to_doc(
        self,
        word_doc: Document,
        block: dict,
        settings: PDFToWordSettings
    ):
        """Add a text block from PDF to Word document."""
        for line in block.get("lines", []):
            para = word_doc.add_paragraph()
            
            for span in line.get("spans", []):
                text = span.get("text", "")
                if not text:
                    continue
                
                run = para.add_run(text)
                
                if settings.preserve_formatting:
                    # Apply font size
                    size = span.get("size", 11)
                    run.font.size = Pt(size)
                    
                    # Apply font name
                    font_name = span.get("font", "")
                    if font_name:
                        # Clean up font name
                        clean_name = font_name.split("+")[-1] if "+" in font_name else font_name
                        try:
                            run.font.name = clean_name
                        except:
                            pass
                    
                    # Check for bold/italic in flags
                    flags = span.get("flags", 0)
                    if flags & 2 ** 0:  # Superscript
                        run.font.superscript = True
                    if flags & 2 ** 1:  # Italic
                        run.font.italic = True
                    if flags & 2 ** 2:  # Serif
                        pass
                    if flags & 2 ** 3:  # Monospace
                        run.font.name = "Courier New"
                    if flags & 2 ** 4:  # Bold
                        run.font.bold = True
                    
                    # Apply color
                    color = span.get("color", 0)
                    if color != 0:
                        r = (color >> 16) & 0xFF
                        g = (color >> 8) & 0xFF
                        b = color & 0xFF
                        run.font.color.rgb = RGBColor(r, g, b)
    
    def _add_image_block_to_doc(
        self,
        word_doc: Document,
        block: dict,
        temp_dir: str
    ):
        """Add an image block from PDF to Word document."""
        try:
            # Get image data
            img_data = block.get("image")
            if not img_data:
                return
            
            # Save image temporarily
            img_ext = block.get("ext", "png")
            img_path = os.path.join(temp_dir, f"img_{id(block)}.{img_ext}")
            
            with open(img_path, "wb") as f:
                f.write(img_data)
            
            # Get dimensions
            width = block.get("width", 0)
            height = block.get("height", 0)
            
            # Add to document (limit width to 6 inches)
            if width > 0:
                word_doc.add_picture(img_path, width=min(Inches(6), Pt(width)))
            else:
                word_doc.add_picture(img_path)
                
        except Exception:
            pass  # Skip failed images
    
    def _add_ocr_text_to_doc(
        self,
        word_doc: Document,
        ocr_data: dict,
        settings: PDFToWordSettings
    ):
        """Add OCR extracted text to Word document."""
        # Group text by line
        n_boxes = len(ocr_data['text'])
        current_line_num = -1
        current_para = None
        
        for i in range(n_boxes):
            text = ocr_data['text'][i]
            conf = int(ocr_data['conf'][i])
            line_num = ocr_data['line_num'][i]
            
            # Skip low confidence or empty text
            if conf < 30 or not text or not text.strip():
                continue
            
            # New line = new paragraph
            if line_num != current_line_num:
                current_para = word_doc.add_paragraph()
                current_line_num = line_num
            
            if current_para:
                # Add text with space
                if current_para.text:
                    current_para.add_run(" ")
                current_para.add_run(text)
    
    def _extract_images_from_page(
        self,
        page: fitz.Page,
        word_doc: Document,
        temp_dir: str
    ):
        """Extract and add images from a PDF page to Word document."""
        try:
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = page.parent.extract_image(xref)
                
                if base_image:
                    img_bytes = base_image["image"]
                    img_ext = base_image["ext"]
                    
                    img_path = os.path.join(
                        temp_dir, 
                        f"page_{page.number}_img_{img_index}.{img_ext}"
                    )
                    
                    with open(img_path, "wb") as f:
                        f.write(img_bytes)
                    
                    # Add to document
                    word_doc.add_picture(img_path, width=Inches(5))
                    
        except Exception:
            pass  # Skip failed image extraction
